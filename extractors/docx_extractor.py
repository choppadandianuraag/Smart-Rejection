"""
DOCX/DOC document extractor.
Handles Word document text extraction.
"""

from pathlib import Path
from typing import Tuple, Dict, Any

from loguru import logger

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import mammoth
    MAMMOTH_AVAILABLE = True
except ImportError:
    MAMMOTH_AVAILABLE = False

from extractors.base import BaseExtractor


class DocxExtractor(BaseExtractor):
    """Extractor for DOCX and DOC documents."""
    
    SUPPORTED_FORMATS = [".docx", ".doc"]
    
    def supports_format(self, file_extension: str) -> bool:
        """Check if Word document format is supported."""
        return file_extension.lower() in self.SUPPORTED_FORMATS
    
    def extract(self, file_path: Path) -> Tuple[str, str, Dict[str, Any]]:
        """
        Extract text from Word document.
        
        Uses python-docx for .docx files and mammoth for HTML/markdown conversion.
        
        Args:
            file_path: Path to Word document
            
        Returns:
            Tuple of (raw_text, markdown_content, metadata)
        """
        metadata = self.get_file_info(file_path)
        metadata["extraction_method"] = None
        
        raw_text = ""
        markdown_content = ""
        
        file_ext = file_path.suffix.lower()
        
        # For DOCX files
        if file_ext == ".docx" and DOCX_AVAILABLE:
            try:
                raw_text = self._extract_with_python_docx(file_path)
                metadata["extraction_method"] = "python-docx"
                logger.info(f"Extracted {len(raw_text)} chars using python-docx")
            except Exception as e:
                logger.warning(f"python-docx extraction failed: {e}")
        
        # Try mammoth for markdown conversion (works for both .doc and .docx)
        if MAMMOTH_AVAILABLE:
            try:
                markdown_content = self._extract_markdown_with_mammoth(file_path)
                if not raw_text:
                    # Use mammoth raw text as fallback
                    raw_text = self._extract_text_with_mammoth(file_path)
                    metadata["extraction_method"] = "mammoth"
                logger.info(f"Converted to markdown using mammoth")
            except Exception as e:
                logger.warning(f"mammoth extraction failed: {e}")
        
        # If no markdown from mammoth, create basic markdown from raw text
        if not markdown_content and raw_text:
            markdown_content = self._text_to_basic_markdown(raw_text)
        
        if not raw_text.strip():
            metadata["requires_ocr"] = True
            logger.warning("Text extraction failed, OCR may be required")
        else:
            metadata["requires_ocr"] = False
        
        return raw_text, markdown_content, metadata
    
    def _extract_with_python_docx(self, file_path: Path) -> str:
        """Extract text using python-docx."""
        doc = Document(file_path)
        paragraphs = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        
        return "\n\n".join(paragraphs)
    
    def _extract_markdown_with_mammoth(self, file_path: Path) -> str:
        """Extract and convert to markdown using mammoth."""
        with open(file_path, "rb") as f:
            result = mammoth.convert_to_markdown(f)
            return result.value
    
    def _extract_text_with_mammoth(self, file_path: Path) -> str:
        """Extract plain text using mammoth."""
        with open(file_path, "rb") as f:
            result = mammoth.extract_raw_text(f)
            return result.value
    
    def _text_to_basic_markdown(self, text: str) -> str:
        """
        Convert plain text to basic markdown format.
        """
        if not text.strip():
            return ""
        
        lines = text.split("\n")
        markdown_lines = []
        
        for line in lines:
            stripped = line.strip()
            if not stripped:
                markdown_lines.append("")
                continue
            
            # Simple heuristics for headings
            if stripped.isupper() and len(stripped) < 50:
                markdown_lines.append(f"## {stripped.title()}")
            else:
                markdown_lines.append(stripped)
        
        return "\n".join(markdown_lines)
